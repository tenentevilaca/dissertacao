"""
Extração de citações e referências usando Claude AI.
Usa leitura XML direta do .docx (mais robusta que python-docx).
"""

import json
import re
import zipfile
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

import anthropic

from modules.dissertation_parser import (
    CitacaoInText, EntradaReferencia,
    _RE_PAREN, _RE_NARRATIVO, _RE_INICIO_REFS, _RE_ENTRADA_REF,
    _normalizar_autores, _extrair_contexto,
)

_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MAX_CHARS = 150_000

_PROMPT = """Você é um especialista em análise de textos acadêmicos brasileiros.

Analise o texto abaixo de uma dissertação/tese acadêmica e realize DUAS tarefas:

TAREFA 1 — Citações no corpo do texto
Identifique TODAS as referências bibliográficas no corpo do texto.
Inclua qualquer formato: (AUTOR, ano), Autor (ano), notas de rodapé [1], citações entre aspas com fonte, etc.

TAREFA 2 — Lista de referências
Extraia CADA ENTRADA da lista de referências (seção "REFERÊNCIAS", "REFERÊNCIAS BIBLIOGRÁFICAS" ou similar no final).

Retorne SOMENTE um JSON, sem texto antes ou depois:
{
  "citacoes": [
    {
      "texto_original": "exatamente como aparece no texto, ex: (CASTELLS, 2001, p. 117)",
      "autores": ["CASTELLS"],
      "ano": "2001",
      "pagina": "117",
      "contexto": "trecho de 200-300 caracteres ao redor da citação no texto"
    }
  ],
  "referencias": [
    {
      "texto_completo": "entrada completa como aparece na lista de referências",
      "autor_sobrenome": "CASTELLS",
      "ano": "2001",
      "titulo": "título identificado da obra"
    }
  ]
}

TEXTO DA DISSERTAÇÃO:
"""


def _texto_de_xml(xml_bytes: bytes) -> str:
    """Extrai texto de um XML do Word, parágrafo por parágrafo."""
    try:
        root = ET.fromstring(xml_bytes)
        linhas = []
        for para in root.iter(f"{{{_NS}}}p"):
            tokens = []
            for t in para.iter(f"{{{_NS}}}t"):
                if t.text:
                    tokens.append(t.text)
            linhas.append("".join(tokens))
        return "\n".join(linhas)
    except Exception:
        return ""


def extrair_texto_docx(caminho: str | Path) -> str:
    """
    Extrai texto completo de um .docx lendo o XML diretamente.
    Captura: corpo, tabelas, cabeçalhos, rodapés e notas de rodapé.
    """
    partes: list[str] = []

    try:
        with zipfile.ZipFile(str(caminho), "r") as z:
            nomes = z.namelist()

            if "word/document.xml" in nomes:
                corpo = _texto_de_xml(z.read("word/document.xml"))
                partes.append(corpo)
            else:
                partes.append("[AVISO: word/document.xml não encontrado no .docx]")

            if "word/footnotes.xml" in nomes:
                partes.append("\n--- NOTAS DE RODAPÉ ---\n")
                partes.append(_texto_de_xml(z.read("word/footnotes.xml")))

            if "word/endnotes.xml" in nomes:
                partes.append("\n--- NOTAS FINAIS ---\n")
                partes.append(_texto_de_xml(z.read("word/endnotes.xml")))

            for nome in nomes:
                if nome.startswith("word/header") or nome.startswith("word/footer"):
                    partes.append(_texto_de_xml(z.read(nome)))

    except zipfile.BadZipFile:
        try:
            import docx as docx_lib
            doc = docx_lib.Document(str(caminho))
            partes.append("\n".join(p.text for p in doc.paragraphs))
        except Exception as e:
            partes.append(f"[ERRO AO LER ARQUIVO: {e}]")
    except Exception as e:
        partes.append(f"[ERRO AO ABRIR .DOCX: {e}]")

    texto = "\n".join(partes)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def extrair_com_regex(texto: str) -> tuple[list[CitacaoInText], list[EntradaReferencia]]:
    """
    Extrai citações e referências usando regex ABNT — rápido e sem custo de API.
    Retorna (citacoes, referencias).
    """
    paragrafos = texto.split("\n")

    # Localiza início da seção de referências
    idx_refs = -1
    for i, p in enumerate(paragrafos):
        if _RE_INICIO_REFS.match(p.strip()):
            idx_refs = i
            break

    # ── Citações (antes das referências) ──────────────────────────────
    limite = idx_refs if idx_refs > 0 else len(paragrafos)
    textos_corpo = paragrafos[:limite]
    citacoes: list[CitacaoInText] = []
    vistos: set[tuple] = set()

    for idx, para in enumerate(textos_corpo):
        contexto = _extrair_contexto(textos_corpo, idx)

        for m in _RE_PAREN.finditer(para):
            autores = _normalizar_autores(m.group(1))
            ano = m.group(2)
            pagina = m.group(3) if m.lastindex and m.lastindex >= 3 else None
            chave_vis = (tuple(autores), ano, pagina, idx)
            if chave_vis not in vistos:
                vistos.add(chave_vis)
                citacoes.append(CitacaoInText(
                    texto_original=m.group(0),
                    autores=autores,
                    ano=ano,
                    pagina=pagina,
                    contexto=contexto,
                    paragrafo_idx=idx,
                ))

        for m in _RE_NARRATIVO.finditer(para):
            autores = _normalizar_autores(m.group(1))
            ano = m.group(2)
            pagina = m.group(3) if m.lastindex and m.lastindex >= 3 else None
            chave_vis = (tuple(autores), ano, pagina, idx)
            if chave_vis not in vistos:
                vistos.add(chave_vis)
                citacoes.append(CitacaoInText(
                    texto_original=m.group(0),
                    autores=autores,
                    ano=ano,
                    pagina=pagina,
                    contexto=contexto,
                    paragrafo_idx=idx,
                ))

    # ── Referências (após cabeçalho REFERÊNCIAS) ───────────────────────
    referencias: list[EntradaReferencia] = []

    if idx_refs >= 0:
        _RE_NOVA_ENTRADA = re.compile(
            r'^[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÇ][A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÇa-záéíóúâêîôûãõàç\'\-]+[,;]'
        )
        buffer: list[str] = []

        def _processar(buf: list[str]) -> Optional[EntradaReferencia]:
            texto_buf = " ".join(buf).strip()
            if not texto_buf:
                return None
            m = _RE_ENTRADA_REF.match(texto_buf)
            if m:
                return EntradaReferencia(
                    texto_completo=texto_buf,
                    autor_sobrenome=m.group(1).strip().upper(),
                    ano=m.group(3).strip(),
                    titulo=m.group(2).strip(),
                )
            m_ano = re.search(r'\b(\d{4})\b', texto_buf)
            m_sob = re.match(
                r'^([A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÇ][A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÇa-záéíóúâêîôûãõàç\'\-]+)',
                texto_buf,
            )
            if m_sob and m_ano:
                return EntradaReferencia(
                    texto_completo=texto_buf,
                    autor_sobrenome=m_sob.group(1).upper(),
                    ano=m_ano.group(1),
                    titulo=texto_buf[:80],
                )
            return None

        for linha in paragrafos[idx_refs + 1:]:
            linha = linha.strip()
            if not linha:
                if buffer:
                    entrada = _processar(buffer)
                    if entrada:
                        referencias.append(entrada)
                    buffer = []
                continue
            if _RE_NOVA_ENTRADA.match(linha) and buffer:
                entrada = _processar(buffer)
                if entrada:
                    referencias.append(entrada)
                buffer = [linha]
            else:
                buffer.append(linha)

        if buffer:
            entrada = _processar(buffer)
            if entrada:
                referencias.append(entrada)

    return citacoes, referencias


class AIExtractor:
    """Usa regex + Claude para extrair citações e referências de qualquer formato acadêmico."""

    def __init__(self, api_key: str, model: str = "claude-sonnet-4-6"):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = model

    def extrair(self, texto: str, log_fn=None) -> dict:
        """Extrai via Claude AI (fallback quando regex não encontra nada)."""
        if not texto.strip():
            if log_fn:
                log_fn("   ✗ Texto vazio — não foi possível ler o arquivo")
            return {"citacoes": [], "referencias": []}

        if log_fn:
            log_fn(f"   Enviando {len(texto):,} caracteres ao Claude…")

        try:
            resposta = self.client.messages.create(
                model=self.model,
                max_tokens=16000,
                messages=[{"role": "user", "content": _PROMPT + texto[:_MAX_CHARS]}],
            )
        except Exception as e:
            if log_fn:
                log_fn(f"   ✗ Erro na API Claude: {e}")
            raise

        raw = resposta.content[0].text.strip()
        stop_reason = resposta.stop_reason

        if log_fn:
            log_fn(f"   Claude respondeu ({len(raw)} chars, stop={stop_reason})")
            log_fn(f"   Prévia: {raw[:200].replace(chr(10),' ')}")
            if stop_reason == "max_tokens":
                log_fn("   ⚠ Resposta cortada! max_tokens atingido — pode faltar parte do JSON")

        # Limpa markdown se presente
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\s*```\s*$", "", raw, flags=re.MULTILINE)

        # Tenta parsear JSON
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            m = re.search(r"\{[\s\S]*\}", raw)
            if m:
                try:
                    return json.loads(m.group())
                except json.JSONDecodeError:
                    pass
            if log_fn:
                log_fn(f"   ✗ JSON inválido. Resposta bruta: {raw[:400]}")
            return {"citacoes": [], "referencias": []}

    def ler_docx(self, caminho: str | Path) -> str:
        """Retorna o texto extraído do .docx."""
        return extrair_texto_docx(caminho)

    def extrair_de_docx(self, caminho: str | Path, log_fn=None) -> dict:
        texto = extrair_texto_docx(caminho)
        if log_fn:
            log_fn(f"   Texto extraído do .docx: {len(texto):,} caracteres")
            if len(texto) < 200:
                log_fn(f"   ⚠ Texto muito curto: {repr(texto[:200])}")
        return self.extrair(texto, log_fn=log_fn)

    def para_objetos(self, dados: dict) -> tuple[list[CitacaoInText], list[EntradaReferencia]]:
        citacoes: list[CitacaoInText] = []
        for i, c in enumerate(dados.get("citacoes", [])):
            try:
                autores = c.get("autores") or ["DESCONHECIDO"]
                if isinstance(autores, str):
                    autores = [autores]
                citacoes.append(CitacaoInText(
                    texto_original=str(c.get("texto_original", "")),
                    autores=[str(a).upper() for a in autores],
                    ano=str(c.get("ano", "s.d.")),
                    pagina=c.get("pagina") or None,
                    contexto=str(c.get("contexto", "")),
                    paragrafo_idx=i,
                ))
            except Exception:
                continue

        referencias: list[EntradaReferencia] = []
        for r in dados.get("referencias", []):
            try:
                referencias.append(EntradaReferencia(
                    texto_completo=str(r.get("texto_completo", "")),
                    autor_sobrenome=str(r.get("autor_sobrenome", "DESCONHECIDO")).upper(),
                    ano=str(r.get("ano", "s.d.")),
                    titulo=str(r.get("titulo", "")),
                ))
            except Exception:
                continue

        return citacoes, referencias

    def cruzar(self, citacoes: list[CitacaoInText], referencias: list[EntradaReferencia]) -> dict:
        idx_refs: dict[str, EntradaReferencia] = {}
        for ref in referencias:
            chave = f"{ref.autor_sobrenome.upper()}_{ref.ano}"
            idx_refs[chave] = ref

        chaves_citadas: set[str] = set()
        citadas_sem_ref: list[CitacaoInText] = []
        pareamentos: list[tuple] = []

        for cit in citacoes:
            encontrou = False
            for autor in cit.autores:
                chave = f"{autor.upper()}_{cit.ano}"
                chaves_citadas.add(chave)
                if chave in idx_refs:
                    pareamentos.append((cit, idx_refs[chave]))
                    encontrou = True
                    break
            if not encontrou:
                citadas_sem_ref.append(cit)

        refs_sem_citacao = [r for k, r in idx_refs.items() if k not in chaves_citadas]

        return {
            "citadas_sem_referencia": citadas_sem_ref,
            "referenciadas_sem_citacao": refs_sem_citacao,
            "pareamentos": pareamentos,
        }


_PROMPT = """Você é um especialista em análise de textos acadêmicos brasileiros.

Analise o texto abaixo de uma dissertação/tese acadêmica e realize DUAS tarefas:

TAREFA 1 — Citações no corpo do texto
Identifique TODAS as referências bibliográficas no corpo do texto.
Inclua qualquer formato: (AUTOR, ano), Autor (ano), notas de rodapé [1], citações entre aspas com fonte, etc.

TAREFA 2 — Lista de referências
Extraia CADA ENTRADA da lista de referências (seção "REFERÊNCIAS", "REFERÊNCIAS BIBLIOGRÁFICAS" ou similar no final).

Retorne SOMENTE um JSON, sem texto antes ou depois:
{
  "citacoes": [
    {
      "texto_original": "exatamente como aparece no texto, ex: (CASTELLS, 2001, p. 117)",
      "autores": ["CASTELLS"],
      "ano": "2001",
      "pagina": "117",
      "contexto": "trecho de 200-300 caracteres ao redor da citação no texto"
    }
  ],
  "referencias": [
    {
      "texto_completo": "entrada completa como aparece na lista de referências",
      "autor_sobrenome": "CASTELLS",
      "ano": "2001",
      "titulo": "título identificado da obra"
    }
  ]
}

TEXTO DA DISSERTAÇÃO:
"""


def _texto_de_xml(xml_bytes: bytes) -> str:
    """Extrai texto de um XML do Word, parágrafo por parágrafo."""
    try:
        root = ET.fromstring(xml_bytes)
        linhas = []
        for para in root.iter(f"{{{_NS}}}p"):
            tokens = []
            for t in para.iter(f"{{{_NS}}}t"):
                if t.text:
                    tokens.append(t.text)
            linhas.append("".join(tokens))
        return "\n".join(linhas)
    except Exception:
        return ""


def extrair_texto_docx(caminho: str | Path) -> str:
    """
    Extrai texto completo de um .docx lendo o XML diretamente.
    Captura: corpo, tabelas, cabeçalhos, rodapés e notas de rodapé.
    """
    partes: list[str] = []

    try:
        with zipfile.ZipFile(str(caminho), "r") as z:
            nomes = z.namelist()

            if "word/document.xml" in nomes:
                corpo = _texto_de_xml(z.read("word/document.xml"))
                partes.append(corpo)
            else:
                partes.append("[AVISO: word/document.xml não encontrado no .docx]")

            if "word/footnotes.xml" in nomes:
                partes.append("\n--- NOTAS DE RODAPÉ ---\n")
                partes.append(_texto_de_xml(z.read("word/footnotes.xml")))

            if "word/endnotes.xml" in nomes:
                partes.append("\n--- NOTAS FINAIS ---\n")
                partes.append(_texto_de_xml(z.read("word/endnotes.xml")))

            for nome in nomes:
                if nome.startswith("word/header") or nome.startswith("word/footer"):
                    partes.append(_texto_de_xml(z.read(nome)))

    except zipfile.BadZipFile:
        try:
            import docx as docx_lib
            doc = docx_lib.Document(str(caminho))
            partes.append("\n".join(p.text for p in doc.paragraphs))
        except Exception as e:
            partes.append(f"[ERRO AO LER ARQUIVO: {e}]")
    except Exception as e:
        partes.append(f"[ERRO AO ABRIR .DOCX: {e}]")

    texto = "\n".join(partes)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


class AIExtractor:
    """Usa Claude para extrair citações e referências de qualquer formato acadêmico."""

    def __init__(self, api_key: str, model: str = "claude-sonnet-4-6"):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = model

    def extrair(self, texto: str, log_fn=None) -> dict:
        if not texto.strip():
            if log_fn:
                log_fn("   ✗ Texto vazio — não foi possível ler o arquivo")
            return {"citacoes": [], "referencias": []}

        if log_fn:
            log_fn(f"   Enviando {len(texto):,} caracteres ao Claude…")

        try:
            resposta = self.client.messages.create(
                model=self.model,
                max_tokens=16000,
                messages=[{"role": "user", "content": _PROMPT + texto[:_MAX_CHARS]}],
            )
        except Exception as e:
            if log_fn:
                log_fn(f"   ✗ Erro na API Claude: {e}")
            raise

        raw = resposta.content[0].text.strip()
        stop_reason = resposta.stop_reason

        if log_fn:
            log_fn(f"   Claude respondeu ({len(raw)} chars, stop={stop_reason})")
            log_fn(f"   Prévia: {raw[:200].replace(chr(10),' ')}")
            if stop_reason == "max_tokens":
                log_fn("   ⚠ Resposta cortada! max_tokens atingido — pode faltar parte do JSON")

        # Limpa markdown se presente
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\s*```\s*$", "", raw, flags=re.MULTILINE)

        # Tenta parsear JSON
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            m = re.search(r"\{[\s\S]*\}", raw)
            if m:
                try:
                    return json.loads(m.group())
                except json.JSONDecodeError:
                    pass
            if log_fn:
                log_fn(f"   ✗ JSON inválido. Resposta bruta: {raw[:400]}")
            return {"citacoes": [], "referencias": []}

    def ler_docx(self, caminho: str | Path) -> str:
        """Retorna o texto extraído do .docx (para uso externo antes de extrair)."""
        return extrair_texto_docx(caminho)

    def extrair_de_docx(self, caminho: str | Path, log_fn=None) -> dict:
        texto = extrair_texto_docx(caminho)
        if log_fn:
            log_fn(f"   Texto extraído do .docx: {len(texto):,} caracteres")
            if len(texto) < 200:
                log_fn(f"   ⚠ Texto muito curto: {repr(texto[:200])}")
        return self.extrair(texto, log_fn=log_fn)

    def para_objetos(self, dados: dict) -> tuple[list[CitacaoInText], list[EntradaReferencia]]:
        citacoes: list[CitacaoInText] = []
        for i, c in enumerate(dados.get("citacoes", [])):
            try:
                autores = c.get("autores") or ["DESCONHECIDO"]
                if isinstance(autores, str):
                    autores = [autores]
                citacoes.append(CitacaoInText(
                    texto_original=str(c.get("texto_original", "")),
                    autores=[str(a).upper() for a in autores],
                    ano=str(c.get("ano", "s.d.")),
                    pagina=c.get("pagina") or None,
                    contexto=str(c.get("contexto", "")),
                    paragrafo_idx=i,
                ))
            except Exception:
                continue

        referencias: list[EntradaReferencia] = []
        for r in dados.get("referencias", []):
            try:
                referencias.append(EntradaReferencia(
                    texto_completo=str(r.get("texto_completo", "")),
                    autor_sobrenome=str(r.get("autor_sobrenome", "DESCONHECIDO")).upper(),
                    ano=str(r.get("ano", "s.d.")),
                    titulo=str(r.get("titulo", "")),
                ))
            except Exception:
                continue

        return citacoes, referencias

    def cruzar(self, citacoes: list[CitacaoInText], referencias: list[EntradaReferencia]) -> dict:
        idx_refs: dict[str, EntradaReferencia] = {}
        for ref in referencias:
            chave = f"{ref.autor_sobrenome.upper()}_{ref.ano}"
            idx_refs[chave] = ref

        chaves_citadas: set[str] = set()
        citadas_sem_ref: list[CitacaoInText] = []
        pareamentos: list[tuple] = []

        for cit in citacoes:
            encontrou = False
            for autor in cit.autores:
                chave = f"{autor.upper()}_{cit.ano}"
                chaves_citadas.add(chave)
                if chave in idx_refs:
                    pareamentos.append((cit, idx_refs[chave]))
                    encontrou = True
                    break
            if not encontrou:
                citadas_sem_ref.append(cit)

        refs_sem_citacao = [r for k, r in idx_refs.items() if k not in chaves_citadas]

        return {
            "citadas_sem_referencia": citadas_sem_ref,
            "referenciadas_sem_citacao": refs_sem_citacao,
            "pareamentos": pareamentos,
        }
