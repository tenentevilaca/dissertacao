"""
Analisador de Dissertação — análise crítica estrutural baseada no template genérico.
Reusa ler_docx, ler_arquivo e _extrair_trecho_relevante do verificador.py.
"""

import io
import json
import re
import zipfile
from pathlib import Path
from typing import Optional

import anthropic

# ── Importações do módulo irmão ──────────────────────────────────────────
from verificador import ler_docx, ler_arquivo, ler_docx_paragrafos, _extrair_trecho_relevante, _STOPWORDS_PT

_MAX_CHARS_CAP   = 14_000
_MAX_CHARS_META  = 7_000
_MAX_CHARS_APOIO = 5_000


# ════════════════════════════════════════════════════════════════════════
# DOWNLOAD DO GOOGLE DRIVE
# ════════════════════════════════════════════════════════════════════════

def _extrair_id_drive(url: str) -> str:
    """Extrai o file ID de qualquer formato de URL do Google Drive."""
    # /file/d/ID/view  ou  /d/ID/
    m = re.search(r'/d/([a-zA-Z0-9_-]{10,})', url)
    if m:
        return m.group(1)
    # ?id=ID  ou  &id=ID
    m = re.search(r'[?&]id=([a-zA-Z0-9_-]{10,})', url)
    if m:
        return m.group(1)
    raise ValueError(
        "Não foi possível identificar o ID do arquivo na URL. "
        "Certifique-se de usar o link de compartilhamento do Google Drive."
    )


def baixar_google_drive_bytes(url: str, log_fn=None) -> tuple:
    """
    Baixa arquivo do Google Drive direto na memória (sem salvar no disco).
    Retorna (bytes, extensao).
    """
    import urllib.request
    import http.cookiejar

    file_id = _extrair_id_drive(url)
    if log_fn:
        log_fn(f"   ID do arquivo: {file_id}")

    jar = http.cookiejar.CookieJar()
    opener = urllib.request.build_opener(urllib.request.HTTPCookieProcessor(jar))
    opener.addheaders = [("User-Agent", "Mozilla/5.0")]

    urls_tentar = [
        f"https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t",
        f"https://drive.google.com/uc?export=download&id={file_id}&confirm=t",
    ]

    data = None
    content_type = ""
    resp_final = None

    for tentativa_url in urls_tentar:
        if log_fn:
            log_fn("   Conectando ao Google Drive...")
        try:
            resp = opener.open(tentativa_url, timeout=300)
            content_type = resp.headers.get("Content-Type", "")
            data = resp.read()
            resp_final = resp
            if "text/html" in content_type and len(data) < 500_000:
                if log_fn:
                    log_fn("   Resposta HTML, tentando alternativa...")
                data = None
                continue
            break
        except Exception as e:
            if log_fn:
                log_fn(f"   Tentativa falhou: {e}")
            continue

    if not data or len(data) < 1000:
        raise ValueError("Nao foi possivel baixar o arquivo. Verifique se o link esta publico.")

    if log_fn:
        log_fn(f"   Download concluido: {len(data)/(1024*1024):.1f} MB")

    # Detecta extensão
    ext = ".zip"
    if resp_final:
        cd = resp_final.headers.get("Content-Disposition", "")
        m_cd = re.search(r'filename[^;=\n]*=\s*["\']?([^"\'\n;]+)', cd)
        if m_cd:
            ext = Path(m_cd.group(1).strip()).suffix or ext
    if "pdf" in content_type:
        ext = ".pdf"
    elif "msword" in content_type or "wordprocessing" in content_type:
        ext = ".docx"

    return data, ext


def baixar_google_drive(url: str, destino: Path, log_fn=None) -> Path:
    """Compat: baixa e salva no disco. Retorna Path."""
    data, ext = baixar_google_drive_bytes(url, log_fn)
    caminho = destino.with_suffix(ext)
    caminho.write_bytes(data)
    return caminho


# ════════════════════════════════════════════════════════════════════════
# PROMPTS
# ════════════════════════════════════════════════════════════════════════

_PROMPT_META = """Você é um assistente especializado em análise de textos acadêmicos brasileiros.

Leia o início desta dissertação e extraia as informações bibliográficas e estruturais.

Retorne SOMENTE um JSON válido (sem markdown, sem texto antes ou depois):
{
  "autor": "Nome completo do autor ou null",
  "titulo": "Título completo da dissertação ou null",
  "programa": "Nome do programa de pós-graduação ou null",
  "instituicao": "Nome da instituição ou null",
  "orientador": "Nome do orientador ou null",
  "area": "Área de concentração ou null",
  "ano": "Ano de conclusão/defesa ou null",
  "problema_pesquisa": "O problema central da pesquisa em 1-2 frases ou null",
  "hipotese": "A hipótese ou tese central em 1-2 frases ou null",
  "objetivos": ["objetivo específico 1", "objetivo específico 2"],
  "capitulos_detectados": ["Título Cap 1", "Título Cap 2"]
}

INÍCIO DA DISSERTAÇÃO:
"""

_PROMPT_ANALISE = """Você é um orientador acadêmico experiente analisando criticamente uma dissertação brasileira.

METADADOS DA DISSERTAÇÃO:
{metadados}

CAPÍTULO EM ANÁLISE: «{titulo_cap}»

TEXTO DO CAPÍTULO (até {n_chars} caracteres):
\"\"\"
{texto_cap}
\"\"\"

Realize a análise crítica completa conforme o framework abaixo.
Retorne SOMENTE um JSON válido (sem markdown, sem texto antes ou depois):
{{
  "titulo_capitulo": "{titulo_cap}",
  "problemas_estruturais": [
    {{
      "descricao": "descrição objetiva do problema",
      "gravidade": "ALTA",
      "localizacao": "onde ocorre (seção, início, fim, ao longo do texto)"
    }}
  ],
  "adequacao_teorica": {{
    "avaliacao": "ADEQUADA",
    "pontos_positivos": ["ponto 1", "ponto 2"],
    "lacunas": ["lacuna 1"]
  }},
  "integracao_empirica": {{
    "avaliacao": "INTEGRADA",
    "observacoes": "descrição de como teoria e dados se articulam (ou não)"
  }},
  "conformidade_abnt": {{
    "problemas_encontrados": ["descrição de problema ABNT"],
    "status": "OK"
  }},
  "estilo_escrita": {{
    "paragrafos_que_iniciam_com_autor": 0,
    "uso_de_travessao": false,
    "uso_excessivo_dois_pontos": false,
    "transicoes_entre_secoes": "PRESENTES",
    "listas_desnecessarias": false,
    "observacoes": "observações gerais sobre estilo"
  }},
  "recomendacoes_priorizadas": [
    "1. [ALTA] recomendação prioritária",
    "2. [MEDIA] recomendação secundária"
  ],
  "pontuacao_geral": "APROVADO_COM_RESSALVAS",
  "resumo_executivo": "Em 3-4 frases: o que está bem e o que precisa de atenção."
}}

Onde "avaliacao"/"status"/"pontuacao_geral" usam apenas: APROVADO | APROVADO_COM_RESSALVAS | REQUER_REVISAO | ADEQUADA | INSUFICIENTE | AUSENTE | INTEGRADA | PARCIAL | OK | COM_PROBLEMAS
"""

_PROMPT_COESAO = """Você é um orientador acadêmico revisando a COESÃO INTERNA de uma dissertação brasileira —
ou seja, como as ideias se encadeiam do geral para o específico e se há repetições/redundâncias entre o capítulo
atual e o restante do texto.

RESUMO DOS DEMAIS CAPÍTULOS (título + trecho inicial de cada um):
\"\"\"
{resumo_outros}
\"\"\"

CAPÍTULO EM ANÁLISE: «{titulo_cap}»

TEXTO DO CAPÍTULO (até {n_chars} caracteres):
\"\"\"
{texto_cap}
\"\"\"

Avalie:
(a) REPETIÇÕES — ideias, definições, citações ou explicações que já aparecem em outro capítulo e são repetidas aqui sem necessidade.
(b) FLUXO GERAL→ESPECÍFICO — o capítulo introduz conceitos gerais antes dos específicos, na ordem adequada para sua posição na dissertação?
(c) ENCADEAMENTO — o capítulo se conecta com naturalidade ao que veio antes e prepara o que vem depois?

Retorne SOMENTE um JSON válido (sem markdown, sem texto antes ou depois):
{{
  "redundancias": [
    {{
      "trecho_atual": "trecho deste capítulo que repete algo (cite diretamente, até 300 chars)",
      "capitulo_relacionado": "título do capítulo onde a ideia já aparece",
      "sugestao": "manter onde, remover ou resumir aqui, com breve justificativa"
    }}
  ],
  "fluxo_geral_especifico": "ADEQUADO",
  "observacoes_fluxo": "breve avaliação da progressão geral→específico deste capítulo",
  "encadeamento": "ADEQUADO",
  "observacoes_encadeamento": "breve avaliação da transição com capítulos vizinhos",
  "sintese": "Em 2-3 frases: avaliação geral da coesão deste capítulo com o restante da dissertação."
}}

Onde "fluxo_geral_especifico" e "encadeamento" usam: ADEQUADO | PARCIAL | INADEQUADO

REGRA: não invente conteúdo. Cite apenas trechos que de fato aparecem nos textos fornecidos.
"""


# ════════════════════════════════════════════════════════════════════════
# DETECÇÃO DE CAPÍTULOS
# ════════════════════════════════════════════════════════════════════════

# Padrões de título de capítulo
_RE_CAP = re.compile(
    r"^(?:"
    r"CAP[IÍ]TULO\s+\d+"                    # CAPÍTULO 1
    r"|(?:CHAPTER|PARTE)\s+\d+"              # CHAPTER 1
    r"|\d{1,2}(?:\.\d{1,2})?\s+[A-ZÁÉÍ]"   # 1. Introdução / 2.1 Contexto
    r"|[IVX]{1,5}\s+[A-ZÁÉÍ]"              # IV Metodologia
    r")",
    re.IGNORECASE,
)

# Também captura linhas curtas em maiúsculas (provável título)
_RE_TITULO_CURTO = re.compile(
    r"^[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÇ][A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÇ\s\-\:]{3,60}$"
)

_SECOES_IGNORAR = {
    "REFERÊNCIAS", "REFERENCIAS", "REFERÊNCIAS BIBLIOGRÁFICAS",
    "REFERÊNCIAS BIBLIOGRAFICAS", "SUMÁRIO", "SUMARIO", "ABSTRACT",
    "RESUMO", "LISTA DE FIGURAS", "LISTA DE TABELAS", "LISTA DE ABREVIATURAS",
    "LISTA DE ABREVIATURAS E SIGLAS", "LISTA DE SIGLAS", "LISTA DE QUADROS",
    "LISTA DE GRÁFICOS", "LISTA DE GRAFICOS", "LISTA DE ANEXOS",
    "AGRADECIMENTOS", "DEDICATÓRIA", "EPÍGRAFE", "ÍNDICE",
}


def detectar_capitulos(texto: str) -> list[tuple[str, str]]:
    """
    Detecta capítulos no texto e retorna lista de (titulo, conteudo).
    Estratégia em dois níveis:
    - Padrões numéricos/CAPÍTULO → sempre são cortes
    - Títulos curtos em caps → apenas se distantes 15+ linhas do anterior
    """
    paragrafos = texto.split("\n")
    cortes: list[tuple[int, str]] = []  # (idx_paragrafo, titulo)

    sumario_idx: int | None = None  # marca início de bloco SUMÁRIO/ÍNDICE (TOC)

    for i, p in enumerate(paragrafos):
        stripped = p.strip()
        if not stripped or len(stripped) > 90:
            # parágrafo de prosa real encerra qualquer bloco de sumário em curso
            if stripped and len(stripped) > 90:
                sumario_idx = None
            continue
        upper = stripped.upper()
        if upper in _SECOES_IGNORAR:
            if upper in {"SUMÁRIO", "SUMARIO", "ÍNDICE", "INDICE"}:
                sumario_idx = i
            continue

        e_cap_explicito = bool(_RE_CAP.match(stripped))
        e_titulo_curto  = (not e_cap_explicito
                           and bool(_RE_TITULO_CURTO.match(stripped))
                           and len(stripped.split()) >= 2)

        # Entradas de sumário (TOC): linha termina em número de página
        # (colado ao título ou separado por espaço/tabulação)
        e_entrada_toc = bool(re.search(r"\d{1,4}$", stripped))

        # Bloco de sumário ainda em curso: ignora cortes dentro do TOC,
        # pois suas entradas são referências fora da ordem do corpo
        if sumario_idx is not None and (i - sumario_idx) < 1000:
            if e_cap_explicito or e_titulo_curto or e_entrada_toc:
                continue

        if e_cap_explicito:
            cortes.append((i, stripped))
            sumario_idx = None
        elif e_titulo_curto:
            # Só inclui se distante 15+ linhas do último corte
            if not cortes or (i - cortes[-1][0]) >= 15:
                cortes.append((i, stripped))
                sumario_idx = None

    if not cortes:
        return [("Dissertação Completa", texto[:_MAX_CHARS_CAP])]

    # Mescla cortes seguidos sem conteúdo entre eles (títulos multi-linha)
    cortes_filtrados: list[tuple[int, str]] = []
    for j, (idx, titulo) in enumerate(cortes):
        if not cortes_filtrados:
            cortes_filtrados.append((idx, titulo))
            continue
        anterior_idx, anterior_titulo = cortes_filtrados[-1]
        linhas_entre = [
            paragrafos[k].strip()
            for k in range(anterior_idx + 1, idx)
            if paragrafos[k].strip()
        ]
        if len(linhas_entre) < 3:
            # Provavelmente sub-título ou continuação — agrega ao título anterior
            cortes_filtrados[-1] = (anterior_idx, f"{anterior_titulo} — {titulo}")
        else:
            cortes_filtrados.append((idx, titulo))

    capitulos: list[tuple[str, str]] = []
    for k, (idx, titulo) in enumerate(cortes_filtrados):
        inicio = idx + 1
        fim = cortes_filtrados[k + 1][0] if k + 1 < len(cortes_filtrados) else len(paragrafos)
        conteudo = "\n".join(paragrafos[inicio:fim]).strip()
        if len(conteudo) > 150:
            capitulos.append((titulo, conteudo))

    return capitulos if capitulos else [("Dissertação Completa", texto[:_MAX_CHARS_CAP])]


# Estilos de parágrafo do Word reconhecidos como títulos de capítulo (nível 1-3)
_RE_ESTILO_TITULO = re.compile(r"^(?:heading|título|titulo|ttulo)\s*[123]$", re.IGNORECASE)


def detectar_capitulos_docx(diss_path: str) -> list[tuple[str, str]] | None:
    """
    Detecção de capítulos baseada nos estilos de parágrafo (Heading 1/2/3) do
    .docx — muito mais confiável que heurísticas de texto puro, pois não
    confunde entradas de sumário/TOC com capítulos reais e captura corretamente
    todos os capítulos, inclusive os iniciais (Introdução, Justificativa,
    Referencial Teórico etc.).

    Retorna None se o arquivo não tiver parágrafos com estilo de título
    suficientes (ex.: docx sem estilos definidos), para que o chamador use o
    fallback heurístico baseado em texto.
    """
    paragrafos = ler_docx_paragrafos(diss_path)
    if not paragrafos:
        return None

    cortes: list[tuple[int, str]] = []
    for i, (texto_p, estilo) in enumerate(paragrafos):
        stripped = (texto_p or "").strip()
        if not stripped or len(stripped) > 150:
            continue
        if _RE_ESTILO_TITULO.match(estilo or ""):
            upper = stripped.upper()
            if upper in _SECOES_IGNORAR:
                continue
            cortes.append((i, stripped))

    # Exige um número mínimo de títulos para considerar a estrutura confiável
    if len(cortes) < 3:
        return None

    capitulos: list[tuple[str, str]] = []
    for k, (idx, titulo) in enumerate(cortes):
        inicio = idx + 1
        fim = cortes[k + 1][0] if k + 1 < len(cortes) else len(paragrafos)
        conteudo = "\n".join(
            p for p, _ in paragrafos[inicio:fim] if p.strip()
        ).strip()
        if len(conteudo) > 150:
            capitulos.append((titulo, conteudo))

    return capitulos if capitulos else None


# ════════════════════════════════════════════════════════════════════════
# LOCALIZADOR DE REFERÊNCIAS NO MATERIAL DE APOIO
# ════════════════════════════════════════════════════════════════════════

# Linha de referência ABNT começa com SOBRENOME(S) em maiúsculas seguido de vírgula
_RE_REF_AUTOR = re.compile(r"^([A-ZÀ-ÝÇ][A-ZÀ-ÝÇ\s\-'\.]{1,60}?),")
_RE_ANO_REF = re.compile(r"\b(1[5-9]\d{2}|20\d{2})\b")


def parsear_lista_referencias(texto: str) -> list[str]:
    """
    Quebra um texto de lista de referências (ABNT) em entradas individuais.
    Entradas começam em linhas no formato 'SOBRENOME, Nome...' — linhas
    seguintes sem esse padrão são consideradas continuação (entradas com
    recuo/quebra de linha).
    """
    refs: list[str] = []
    atual = ""
    for linha in texto.split("\n"):
        l = linha.strip()
        if not l:
            continue
        if _RE_REF_AUTOR.match(l):
            if atual:
                refs.append(atual.strip())
            atual = l
        else:
            atual = f"{atual} {l}".strip() if atual else l
    if atual:
        refs.append(atual.strip())
    return refs


def _extrair_autor_ano_referencia(ref: str) -> tuple[str, str]:
    """Extrai o(s) sobrenome(s) do(s) primeiro(s) autor(es) e o ano de uma referência."""
    m = _RE_REF_AUTOR.match(ref.strip())
    autor = m.group(1).strip().rstrip(".") if m else ""
    # Em referências com múltiplos autores, considera só o primeiro sobrenome
    autor = re.split(r"\s*;\s*|\s+E\s+|\s*&\s*", autor)[0].strip()
    anos = _RE_ANO_REF.findall(ref)
    ano = anos[-1] if anos else ""
    return autor, ano


_PROMPT_DESAMBIGUAR_REF = """Você está ajudando a localizar, entre arquivos de apoio de uma dissertação, qual arquivo corresponde a uma determinada referência bibliográfica.

REFERÊNCIA: {referencia}

Abaixo estão trechos de {n} arquivo(s) candidatos (cada um pode ser uma coletânea com vários textos):

{candidatos}

Qual arquivo é (ou contém) a obra citada na referência acima? Responda SOMENTE com JSON:
{{"indice": <número do arquivo correto, 0 se nenhum corresponder>, "justificativa": "uma frase curta"}}
"""


def _desambiguar_com_ia(
    ref: str,
    candidatos_nomes: list[str],
    material: dict[str, str],
    api_key: str,
) -> Optional[str]:
    """Usa IA (modelo barato) para escolher, entre candidatos empatados, qual
    arquivo realmente corresponde à referência. Retorna o nome do arquivo
    escolhido, ou None se nenhum corresponder / em caso de erro."""
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)

    blocos = []
    for i, nome in enumerate(candidatos_nomes, 1):
        trecho = material[nome][:1500]
        blocos.append(f"ARQUIVO {i}: {nome}\n\"\"\"{trecho}\"\"\"")

    prompt = _PROMPT_DESAMBIGUAR_REF.format(
        referencia=ref, n=len(candidatos_nomes), candidatos="\n\n".join(blocos)
    )

    try:
        resultado = _chamar_claude(client, prompt, max_tokens=200, model="claude-haiku-4-5")
        idx = resultado.get("indice", 0)
        if isinstance(idx, int) and 1 <= idx <= len(candidatos_nomes):
            return candidatos_nomes[idx - 1]
    except Exception:
        pass
    return None


def localizar_referencias(
    referencias: list[str],
    material: dict[str, str],
    api_key: str = "",
    log_fn=None,
) -> list[dict]:
    """
    Para cada referência da lista, procura nos textos do material de apoio
    ocorrências do sobrenome do autor (e do ano, quando disponível),
    retornando os arquivos mais prováveis de conter a obra citada.

    Como os arquivos podem ser coletâneas (vários artigos em um único PDF),
    a busca varre o texto inteiro de cada arquivo, não apenas o início.
    """
    materiais_lower = {nome: txt.lower() for nome, txt in material.items()}

    resultados: list[dict] = []
    for ref in referencias:
        autor, ano = _extrair_autor_ano_referencia(ref)
        candidatos: list[tuple[float, str]] = []

        if autor:
            autor_low = autor.lower()
            for nome_arq, texto_low in materiais_lower.items():
                n_autor = texto_low.count(autor_low)
                if n_autor == 0:
                    continue

                # Se o ano da referência é conhecido mas NÃO aparece em
                # lugar nenhum do arquivo, é forte indício de obra errada.
                if ano and ano not in material[nome_arq]:
                    continue

                score = 0.0

                # Sinal forte: nome do autor está no PRÓPRIO nome do arquivo
                if autor_low in Path(nome_arq).stem.lower():
                    score += 50

                # Sinal forte: autor e ano aparecem próximos um do outro
                # (capa/cabeçalho do arquivo, ou citação da própria obra)
                if ano:
                    for m in re.finditer(re.escape(autor_low), texto_low):
                        janela = texto_low[max(0, m.start() - 150): m.end() + 150]
                        if ano in janela:
                            score += 10

                # Sinal fraco: ocorrências avulsas do sobrenome (capado para
                # não deixar sobrenomes comuns dominarem o ranking)
                score += min(n_autor, 5)

                if score >= 10:
                    candidatos.append((score, nome_arq))

        candidatos.sort(key=lambda x: x[0], reverse=True)

        # Mantém só candidatos competitivos: o melhor, e demais apenas se
        # tiverem pontuação próxima do melhor — evita listar arquivos sem
        # relação real com a referência. Score mínimo de 10 já exige um
        # sinal forte (autor no nome do arquivo, ou autor+ano próximos).
        arquivos: list[str] = []
        if candidatos:
            melhor_score = candidatos[0][0]
            limite = melhor_score * 0.6
            for score, nome_arq in candidatos[:3]:
                if score >= limite:
                    arquivos.append(nome_arq)

        # Caso ambíguo (mais de um candidato competitivo) e temos chave de
        # API: pede para a IA (modelo barato) escolher o arquivo correto.
        if api_key.strip() and len(arquivos) > 1:
            escolhido = _desambiguar_com_ia(ref, arquivos, material, api_key)
            if escolhido:
                arquivos = [escolhido]

        # Nenhum sinal forte de nome de arquivo/proximidade autor+ano: o
        # autor e/ou título da obra podem aparecer no CONTEÚDO do arquivo
        # com nome diferente (ex.: "documento1.pdf"). Se há chave de API,
        # monta uma lista de candidatos "fracos" (autor citado em algum
        # lugar e/ou palavras do título presentes) e pede à IA para ler o
        # conteúdo e decidir se algum deles é a obra referenciada.
        if not arquivos and api_key.strip():
            titulo_palavras = [
                p for p in re.findall(r"[a-záéíóúâêîôûãõàç]{4,}", ref.lower())
                if p not in _STOPWORDS_PT
            ][:8]
            fracos: list[tuple[float, str]] = []
            for nome_arq, texto_low in materiais_lower.items():
                pontos = 0.0
                if autor:
                    pontos += texto_low.count(autor.lower()) * 2
                if ano and ano in material[nome_arq]:
                    pontos += 1
                pontos += sum(1 for p in titulo_palavras if p in texto_low[:3000])
                if pontos > 0:
                    fracos.append((pontos, nome_arq))
            fracos.sort(key=lambda x: -x[0])
            candidatos_fracos = [nome for _, nome in fracos[:6]]
            if candidatos_fracos:
                escolhido = _desambiguar_com_ia(ref, candidatos_fracos, material, api_key)
                if escolhido:
                    arquivos = [escolhido]

        resultados.append({
            "referencia": ref,
            "autor": autor or "—",
            "ano": ano or "—",
            "arquivos": arquivos,
            "encontrado": bool(arquivos),
        })

        if log_fn:
            log_fn(len(resultados), len(referencias), ref, bool(arquivos), arquivos)

    return resultados


def extrair_material_de_zip_com_bytes(zip_source) -> dict[str, tuple[str, bytes]]:
    """
    Como `extrair_material_de_zip`, mas também retorna os bytes originais de
    cada arquivo (para permitir download direto pelo nome do arquivo).
    Retorna dict nome_arquivo → (texto, bytes).
    """
    if isinstance(zip_source, (str, Path)):
        zip_bytes = Path(zip_source).read_bytes()
    else:
        zip_bytes = zip_source

    material: dict[str, tuple[str, bytes]] = {}

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
        for info in z.infolist():
            nome = Path(info.filename).name
            if nome.startswith("~$") or info.is_dir():
                continue
            ext = Path(info.filename).suffix.lower()
            if ext not in EXTENSOES_APOIO:
                continue
            try:
                data = z.read(info)
                texto = _ler_bytes(data, ext)
                if texto and len(texto) > 100 and not texto.startswith("[ERRO"):
                    material[nome[:80]] = (texto, data)
            except Exception:
                continue

    return material


# ════════════════════════════════════════════════════════════════════════
# EXTRAÇÃO DE MATERIAL DE APOIO (ZIP)
# ════════════════════════════════════════════════════════════════════════

EXTENSOES_APOIO = {".pdf", ".docx", ".doc", ".txt"}


def _ler_bytes(data: bytes, ext: str) -> str:
    """Lê conteúdo de um arquivo a partir de bytes, sem salvar no disco."""
    try:
        if ext == ".txt":
            return data.decode("utf-8", errors="replace")
        elif ext in (".docx", ".doc"):
            from verificador import _xml_para_texto
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                partes = []
                if "word/document.xml" in z.namelist():
                    partes.append(_xml_para_texto(z.read("word/document.xml")))
                if "word/footnotes.xml" in z.namelist():
                    partes.append(_xml_para_texto(z.read("word/footnotes.xml")))
                return "\n".join(partes)
        elif ext == ".pdf":
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            return "\n".join(page.get_text() for page in doc)
    except Exception as e:
        return f"[ERRO ao ler: {e}]"
    return ""


def extrair_material_de_zip(zip_source, destino: Path = None) -> dict[str, str]:
    """
    Lê arquivos suportados de um ZIP diretamente na memória.
    zip_source pode ser Path (arquivo no disco) ou bytes.
    Retorna dict nome_arquivo → texto.
    """
    if isinstance(zip_source, (str, Path)):
        zip_bytes = Path(zip_source).read_bytes()
    else:
        zip_bytes = zip_source

    material: dict[str, str] = {}

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
        for info in z.infolist():
            nome = Path(info.filename).name
            if nome.startswith("~$") or info.is_dir():
                continue
            ext = Path(info.filename).suffix.lower()
            if ext not in EXTENSOES_APOIO:
                continue
            try:
                data = z.read(info)
                texto = _ler_bytes(data, ext)
                if texto and len(texto) > 100 and not texto.startswith("[ERRO"):
                    material[nome[:80]] = texto
            except Exception:
                continue

    return material


def carregar_material_de_pasta(pasta: Path) -> dict[str, str]:
    """
    Lê arquivos de uma pasta já extraída (quando vieram como uploads individuais).
    """
    material: dict[str, str] = {}
    for arq in pasta.rglob("*"):
        if arq.is_file() and arq.suffix.lower() in EXTENSOES_APOIO:
            if arq.name.startswith("~$"):
                continue
            texto = ler_arquivo(arq)
            if texto and len(texto) > 100 and not texto.startswith("[ERRO"):
                material[arq.name] = texto
    return material


# ════════════════════════════════════════════════════════════════════════
# CHAMADAS AO CLAUDE
# ════════════════════════════════════════════════════════════════════════

def _chamar_claude(
    client: anthropic.Anthropic,
    prompt: str,
    max_tokens: int = 4096,
    model: str = "claude-sonnet-4-6",
) -> dict:
    """Chama o Claude e parseia o JSON retornado."""
    resp = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = resp.content[0].text.strip()
    # Remove markdown se presente
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"\s*```\s*$", "", raw, flags=re.MULTILINE)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{[\s\S]*\}", raw)
        if m:
            try:
                return json.loads(m.group())
            except json.JSONDecodeError:
                pass
        return {"_erro_parse": raw[:500]}


def extrair_metadados(texto: str, api_key: str) -> dict:
    client = anthropic.Anthropic(api_key=api_key)
    prompt = _PROMPT_META + texto[:_MAX_CHARS_META]
    return _chamar_claude(client, prompt, max_tokens=2048)


def analisar_capitulo(
    titulo_cap: str,
    texto_cap: str,
    metadados: dict,
    api_key: str,
) -> dict:
    client = anthropic.Anthropic(api_key=api_key)
    meta_str = json.dumps(metadados, ensure_ascii=False, indent=2)
    trecho = texto_cap[:_MAX_CHARS_CAP]
    prompt = _PROMPT_ANALISE.format(
        metadados=meta_str[:1000],
        titulo_cap=titulo_cap,
        texto_cap=trecho,
        n_chars=len(trecho),
    )
    return _chamar_claude(client, prompt, max_tokens=3000)


_MAX_CHARS_COESAO = 4000


def _resumo_outros_capitulos(capitulos: list[tuple[str, str]], idx_atual: int, max_chars_cada: int = 600) -> str:
    blocos = []
    for i, (titulo, conteudo) in enumerate(capitulos):
        if i == idx_atual:
            continue
        blocos.append(f"[{titulo}]\n{conteudo[:max_chars_cada]}")
    return "\n\n---\n\n".join(blocos)


def analisar_coesao_capitulo(
    titulo_cap: str,
    texto_cap: str,
    resumo_outros: str,
    api_key: str,
) -> dict:
    if not resumo_outros.strip():
        return {"redundancias": [], "fluxo_geral_especifico": "ADEQUADO",
                "observacoes_fluxo": "", "encadeamento": "ADEQUADO",
                "observacoes_encadeamento": "",
                "sintese": "Capítulo único — não há outros capítulos para comparar."}

    client = anthropic.Anthropic(api_key=api_key)
    prompt = _PROMPT_COESAO.format(
        resumo_outros=resumo_outros[:_MAX_CHARS_COESAO * 2],
        titulo_cap=titulo_cap,
        texto_cap=texto_cap[:_MAX_CHARS_COESAO],
        n_chars=_MAX_CHARS_COESAO,
    )
    return _chamar_claude(client, prompt, max_tokens=2000)


# ════════════════════════════════════════════════════════════════════════
# ORQUESTRADOR PRINCIPAL
# ════════════════════════════════════════════════════════════════════════

def analisar_dissertacao(
    diss_path: str,
    api_key: str,
    log_fn,
) -> dict:
    """
    Pipeline completo:
      1. Lê dissertação
      2. Extrai metadados via IA
      3. Detecta capítulos
      4. Analisa cada capítulo via IA (qualidade/estrutura)
      5. Analisa coesão interna: redundâncias e fluxo geral→específico
    """

    # ── Etapa 1: Leitura ────────────────────────────────────────────────
    log_fn("📄 ETAPA 1: Lendo a dissertação…", "etapa")
    texto = ler_docx(diss_path)
    n = len(texto)
    log_fn(f"   {n:,} caracteres extraídos")
    if n < 300:
        log_fn("   ⚠ Texto muito curto — o arquivo pode estar vazio ou corrompido", "warn")

    # ── Etapa 2: Metadados ──────────────────────────────────────────────
    log_fn("🔍 ETAPA 2: Extraindo metadados via IA…", "etapa")
    try:
        metadados = extrair_metadados(texto, api_key)
        autor = metadados.get("autor") or "Não identificado"
        titulo = metadados.get("titulo") or "Não identificado"
        log_fn(f"   Autor : {autor}")
        log_fn(f"   Título: {titulo}")
        log_fn(f"   Instituição: {metadados.get('instituicao') or '—'}")
        log_fn(f"   Capítulos detectados pela IA: {len(metadados.get('capitulos_detectados') or [])}")
    except Exception as e:
        log_fn(f"   ⚠ Erro ao extrair metadados: {e}", "warn")
        metadados = {}

    # ── Etapa 3: Detecção de capítulos ──────────────────────────────────
    log_fn("📑 ETAPA 3: Detectando estrutura de capítulos…", "etapa")
    capitulos = detectar_capitulos_docx(diss_path)
    if capitulos:
        log_fn(f"   (estrutura obtida pelos estilos de título do .docx)")
    else:
        capitulos = detectar_capitulos(texto)
    log_fn(f"   {len(capitulos)} capítulo(s) / seção(ões) detectado(s)")
    for titulo_cap, conteudo in capitulos[:5]:
        log_fn(f"   • {titulo_cap[:60]} ({len(conteudo):,} chars)")

    # ── Etapas 4+5: Análise por capítulo ────────────────────────────────
    resultados_caps: list[dict] = []
    total = len(capitulos)

    for i, (titulo_cap, conteudo) in enumerate(capitulos, 1):
        log_fn(f"🤖 [{i}/{total}] Analisando: «{titulo_cap[:55]}»…", "etapa")

        # Análise crítica
        try:
            analise = analisar_capitulo(titulo_cap, conteudo, metadados, api_key)
            pontuacao = analise.get("pontuacao_geral", "—")
            resumo = analise.get("resumo_executivo", "")
            log_fn(f"   Pontuação: {pontuacao}")
            if resumo:
                log_fn(f"   {resumo[:120]}")
        except Exception as e:
            log_fn(f"   ✗ Erro na análise: {e}", "warn")
            analise = {"titulo_capitulo": titulo_cap, "_erro": str(e)}

        # Análise de coesão interna (redundâncias + fluxo geral→específico)
        coesao: dict = {}
        if total > 1:
            log_fn("   🔎 Verificando coesão com os demais capítulos…")
            try:
                resumo_outros = _resumo_outros_capitulos(capitulos, i - 1)
                coesao = analisar_coesao_capitulo(titulo_cap, conteudo, resumo_outros, api_key)
                n_red = len(coesao.get("redundancias") or [])
                fluxo = coesao.get("fluxo_geral_especifico", "—")
                log_fn(f"   {n_red} redundância(s) | fluxo geral→específico: {fluxo}")
            except Exception as e:
                log_fn(f"   ⚠ Erro na verificação de coesão: {e}", "warn")
                coesao = {"_erro": str(e)}
        else:
            log_fn("   ℹ Capítulo único — verificação de coesão pulada")

        resultados_caps.append({
            "titulo": titulo_cap,
            "n_chars": len(conteudo),
            "analise": analise,
            "coesao": coesao,
        })

    log_fn(f"✅ Análise concluída — {len(resultados_caps)} capítulo(s) processado(s)", "ok")

    return {
        "metadados": metadados,
        "n_capitulos": len(capitulos),
        "capitulos": resultados_caps,
        "_capitulos_raw": capitulos,
    }


# ════════════════════════════════════════════════════════════════════════
# GERAÇÃO DE RELATÓRIO
# ════════════════════════════════════════════════════════════════════════

_COR_PONT = {
    "APROVADO": "#16a34a",
    "APROVADO_COM_RESSALVAS": "#d97706",
    "REQUER_REVISAO": "#dc2626",
}
_ICON_PONT = {
    "APROVADO": "✅",
    "APROVADO_COM_RESSALVAS": "⚠️",
    "REQUER_REVISAO": "❌",
}
_COR_REL = {
    "VALIDA": "#16a34a",
    "CONTRADIZ": "#dc2626",
    "COMPLEMENTA": "#2563eb",
}
_CLASSIF_LABEL = {
    "PLENAMENTE_CONSONANTE": ("✅ Plenamente consonante", "#16a34a"),
    "CONSONANTE_COM_RESSALVA": ("✅⚠️ Consonante com ressalva", "#65a30d"),
    "IMPRECISAO_ATRIBUICAO": ("⚠️ Imprecisão de atribuição", "#d97706"),
    "DISTORCAO_PARCIAL": ("🔶 Distorção parcial", "#ea580c"),
    "DISTORCAO_GRAVE": ("❌ Distorção grave", "#dc2626"),
}


def _esc(s) -> str:
    return str(s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def gerar_relatorio_analise(resultado: dict, rel_dir: Path) -> None:
    from datetime import datetime

    rel_dir = Path(rel_dir)
    rel_dir.mkdir(parents=True, exist_ok=True)

    meta = resultado.get("metadados") or {}
    caps = resultado.get("capitulos") or []
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── JSON ─────────────────────────────────────────────────────────────
    resultado_json = {k: v for k, v in resultado.items() if not k.startswith("_")}
    (rel_dir / "relatorio_analise.json").write_text(
        json.dumps(resultado_json, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # ── Estatísticas gerais ──────────────────────────────────────────────
    pontuacoes = [c["analise"].get("pontuacao_geral", "") for c in caps]
    n_aprovado    = sum(1 for p in pontuacoes if p == "APROVADO")
    n_ressalvas   = sum(1 for p in pontuacoes if p == "APROVADO_COM_RESSALVAS")
    n_revisao     = sum(1 for p in pontuacoes if p == "REQUER_REVISAO")
    total_red     = sum(len(c.get("coesao", {}).get("redundancias") or []) for c in caps)

    # ── HTML ─────────────────────────────────────────────────────────────
    autor     = _esc(meta.get("autor") or "Não identificado")
    titulo_d  = _esc(meta.get("titulo") or "Dissertação")
    inst      = _esc(meta.get("instituicao") or "—")
    orientador= _esc(meta.get("orientador") or "—")
    programa  = _esc(meta.get("programa") or "—")
    area      = _esc(meta.get("area") or "—")
    ano       = _esc(meta.get("ano") or "—")

    problema  = _esc(meta.get("problema_pesquisa") or "—")
    hipotese  = _esc(meta.get("hipotese") or "—")
    objetivos = meta.get("objetivos") or []

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Análise — {titulo_d}</title>
<style>
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
:root {{
  --azul: #1a3a5c; --azul2: #2563eb; --verde: #16a34a;
  --amarelo: #d97706; --vermelho: #dc2626; --cinza: #64748b;
  --fundo: #f1f5f9; --branco: #ffffff; --borda: #e2e8f0;
}}
body {{ font-family: 'Segoe UI', system-ui, sans-serif; background: var(--fundo);
        color: #1e293b; line-height: 1.6; }}
header {{ background: var(--azul); color: #fff; padding: 24px 32px; }}
header h1 {{ font-size: 1.3rem; margin-bottom: 4px; }}
header p  {{ font-size: .85rem; opacity: .75; }}
main {{ max-width: 960px; margin: 0 auto; padding: 32px 16px 80px; }}
h2 {{ font-size: 1.05rem; color: var(--azul); margin: 28px 0 14px; }}
h3 {{ font-size: .95rem; color: #334155; margin: 16px 0 8px; }}
.card {{ background: var(--branco); border-radius: 12px; border: 1px solid var(--borda);
         padding: 24px; margin-bottom: 20px; box-shadow: 0 2px 8px rgba(0,0,0,.06); }}
.meta-grid {{ display: grid; grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
              gap: 12px; margin-bottom: 20px; }}
.meta-item {{ background: #f8fafc; border-radius: 8px; padding: 12px;
              border: 1px solid var(--borda); }}
.meta-item label {{ font-size: .72rem; font-weight: 700; color: var(--cinza);
                    text-transform: uppercase; letter-spacing: .05em; display: block; margin-bottom: 4px; }}
.meta-item span  {{ font-size: .9rem; color: #1e293b; font-weight: 600; }}
.stats-row {{ display: flex; gap: 12px; flex-wrap: wrap; margin-bottom: 24px; }}
.stat {{ border-radius: 10px; padding: 14px 20px; text-align: center; min-width: 110px;
         border: 1px solid var(--borda); }}
.stat strong {{ display: block; font-size: 2rem; font-weight: 800; }}
.stat span   {{ font-size: .75rem; color: var(--cinza); }}
.s-verde  {{ background: #f0fdf4; color: var(--verde); }}
.s-amarelo{{ background: #fffbeb; color: var(--amarelo); }}
.s-vermelho{{ background: #fef2f2; color: var(--vermelho); }}
.s-azul   {{ background: #eff6ff; color: var(--azul2); }}
.cap-card {{ background: var(--branco); border-radius: 12px; border: 1px solid var(--borda);
             margin-bottom: 20px; overflow: hidden; }}
.cap-header {{ padding: 16px 20px; cursor: pointer; display: flex; justify-content: space-between;
               align-items: center; user-select: none; }}
.cap-header:hover {{ background: #f8fafc; }}
.cap-title  {{ font-weight: 700; font-size: .95rem; color: var(--azul); }}
.badge {{ display: inline-block; padding: 3px 10px; border-radius: 20px;
          font-size: .72rem; font-weight: 700; letter-spacing: .04em; }}
.badge-verde    {{ background: #dcfce7; color: var(--verde); }}
.badge-amarelo  {{ background: #fef9c3; color: #92400e; }}
.badge-vermelho {{ background: #fee2e2; color: var(--vermelho); }}
.badge-azul     {{ background: #dbeafe; color: var(--azul2); }}
.badge-cinza    {{ background: #f1f5f9; color: var(--cinza); }}
.cap-body {{ padding: 0 20px 20px; border-top: 1px solid var(--borda); display: none; }}
.cap-body.aberto {{ display: block; }}
.bloco {{ margin-top: 16px; }}
.bloco-titulo {{ font-size: .78rem; font-weight: 700; color: var(--cinza);
                 text-transform: uppercase; letter-spacing: .05em; margin-bottom: 8px; }}
ul.lista {{ padding-left: 18px; }}
ul.lista li {{ margin-bottom: 4px; font-size: .88rem; }}
.par-item {{ background: #f8fafc; border-radius: 8px; padding: 12px; margin-bottom: 8px;
             border-left: 4px solid var(--borda); }}
.par-item.valida     {{ border-left-color: var(--verde); }}
.par-item.contradiz  {{ border-left-color: var(--vermelho); }}
.par-item.complementa{{ border-left-color: var(--azul2); }}
.par-item .rel-tag {{ font-size: .7rem; font-weight: 700; padding: 2px 8px;
                      border-radius: 10px; display: inline-block; margin-bottom: 6px; }}
.par-item .argumento {{ font-size: .85rem; color: #334155; margin-bottom: 6px; }}
.par-item .trecho    {{ font-size: .82rem; color: #64748b; font-style: italic;
                        border-left: 2px solid #cbd5e1; padding-left: 8px; }}
.par-item .fonte     {{ font-size: .72rem; color: var(--cinza); margin-top: 6px; }}
.sintese {{ background: #f0f9ff; border-radius: 8px; padding: 14px;
            border: 1px solid #bae6fd; font-size: .88rem; color: #0369a1; margin-top: 12px; }}
.resumo-exec {{ background: #f0fdf4; border-radius: 8px; padding: 14px;
                border: 1px solid #bbf7d0; font-size: .88rem; color: #166534; margin-bottom: 12px; }}
.prob-item {{ background: #fef2f2; border-radius: 6px; padding: 10px 12px;
              margin-bottom: 6px; font-size: .85rem; border-left: 3px solid #fca5a5; }}
.prob-item.media  {{ background: #fffbeb; border-left-color: #fcd34d; }}
.prob-item.baixa  {{ background: #f8fafc; border-left-color: #cbd5e1; }}
.prob-loc {{ font-size: .72rem; color: var(--cinza); margin-top: 4px; }}
table.abnt {{ width: 100%; border-collapse: collapse; font-size: .84rem; margin-top: 8px; }}
table.abnt td {{ padding: 6px 10px; border-bottom: 1px solid var(--borda); }}
table.abnt tr:last-child td {{ border-bottom: none; }}
.check {{ color: var(--verde); font-weight: 700; }}
.cross {{ color: var(--vermelho); font-weight: 700; }}
.problema-pesquisa {{ background: #eff6ff; border-radius: 8px; padding: 14px;
                      border: 1px solid #bfdbfe; font-size: .9rem; margin-bottom: 12px; }}
.chip {{ display: inline-block; background: #dbeafe; color: var(--azul2);
         border-radius: 20px; padding: 3px 12px; font-size: .78rem;
         font-weight: 600; margin: 3px 3px 0 0; }}
footer {{ text-align: center; font-size: .78rem; color: var(--cinza); margin-top: 40px; padding: 20px; }}
@media (max-width: 600px) {{
  .meta-grid {{ grid-template-columns: 1fr 1fr; }}
  .stats-row {{ gap: 8px; }}
  .stat {{ min-width: 80px; padding: 10px; }}
}}
</style>
</head>
<body>
<header>
  <h1>📋 Análise Crítica de Dissertação</h1>
  <p>Gerado em {_esc(ts)} · Framework de Análise Genérico v1.0</p>
</header>
<main>

<!-- ── Metadados ──────────────────────────────────────────────────────── -->
<div class="card">
  <h2>📌 Identificação da Obra</h2>
  <div class="meta-grid">
    <div class="meta-item"><label>Autor</label><span>{autor}</span></div>
    <div class="meta-item"><label>Instituição</label><span>{inst}</span></div>
    <div class="meta-item"><label>Programa</label><span>{programa}</span></div>
    <div class="meta-item"><label>Orientador</label><span>{orientador}</span></div>
    <div class="meta-item"><label>Área</label><span>{area}</span></div>
    <div class="meta-item"><label>Ano</label><span>{ano}</span></div>
  </div>
  <div style="background:#f8fafc;border-radius:8px;padding:14px;border:1px solid var(--borda)">
    <b style="font-size:.92rem;">{titulo_d}</b>
  </div>

  {"<div class='problema-pesquisa' style='margin-top:12px'><b>Problema de pesquisa:</b><br>" + problema + "</div>" if meta.get("problema_pesquisa") else ""}
  {"<div class='problema-pesquisa' style='background:#f0fdf4;border-color:#bbf7d0;margin-top:8px'><b>Hipótese central:</b><br>" + hipotese + "</div>" if meta.get("hipotese") else ""}

  {('<div style="margin-top:12px"><b style="font-size:.8rem;color:var(--cinza)">OBJETIVOS ESPECÍFICOS</b><br>'
    + "".join(f'<span class="chip">{_esc(o)}</span>' for o in objetivos)
    + '</div>') if objetivos else ""}
</div>

<!-- ── Painel de resultados ────────────────────────────────────────────── -->
<div class="card">
  <h2>📊 Resumo da Análise</h2>
  <div class="stats-row">
    <div class="stat s-verde"><strong>{n_aprovado}</strong><span>Aprovados</span></div>
    <div class="stat s-amarelo"><strong>{n_ressalvas}</strong><span>Com ressalvas</span></div>
    <div class="stat s-vermelho"><strong>{n_revisao}</strong><span>Requer revisão</span></div>
    <div class="stat s-azul"><strong>{total_red}</strong><span>Redundâncias</span></div>
    <div class="stat" style="background:#f8fafc;color:var(--cinza)"><strong>{len(caps)}</strong><span>Capítulos</span></div>
  </div>
</div>

<!-- ── Capítulos ──────────────────────────────────────────────────────── -->
<h2>📖 Análise por Capítulo</h2>
<p style="font-size:.84rem;color:var(--cinza);margin-bottom:16px;">Clique em cada capítulo para expandir os detalhes.</p>
"""

    # ── Cards por capítulo ───────────────────────────────────────────────
    for i, cap in enumerate(caps):
        titulo_cap = _esc(cap.get("titulo", f"Capítulo {i+1}"))
        analise = cap.get("analise") or {}
        coesao  = cap.get("coesao") or {}

        pont = analise.get("pontuacao_geral", "")
        cor_pont = _COR_PONT.get(pont, "#64748b")
        badge_cls = {"APROVADO": "badge-verde", "APROVADO_COM_RESSALVAS": "badge-amarelo",
                     "REQUER_REVISAO": "badge-vermelho"}.get(pont, "badge-cinza")
        icon_pont = _ICON_PONT.get(pont, "•")

        resumo_exec = _esc(analise.get("resumo_executivo") or "")

        # Problemas estruturais
        probs = analise.get("problemas_estruturais") or []
        probs_html = ""
        for prob in probs:
            grav = str(prob.get("gravidade", "")).upper()
            cls_grav = {"MEDIA": "media", "BAIXA": "baixa"}.get(grav, "")
            loc = _esc(prob.get("localizacao") or "")
            probs_html += f"""<div class="prob-item {cls_grav}">
              <b>[{_esc(grav)}]</b> {_esc(prob.get("descricao", ""))}
              {f'<div class="prob-loc">📍 {loc}</div>' if loc else ""}
            </div>"""

        # Adequação teórica
        teo = analise.get("adequacao_teorica") or {}
        teo_av = _esc(teo.get("avaliacao", "—"))
        teo_pos = teo.get("pontos_positivos") or []
        teo_lac = teo.get("lacunas") or []

        # Integração empírica
        emp = analise.get("integracao_empirica") or {}
        emp_av = _esc(emp.get("avaliacao", "—"))
        emp_obs = _esc(emp.get("observacoes") or "")

        # ABNT
        abnt = analise.get("conformidade_abnt") or {}
        abnt_probs = abnt.get("problemas_encontrados") or []
        abnt_status = _esc(abnt.get("status") or "")

        # Estilo
        estilo = analise.get("estilo_escrita") or {}
        est_autor = estilo.get("paragrafos_que_iniciam_com_autor", 0)
        est_trav  = estilo.get("uso_de_travessao", False)
        est_2pts  = estilo.get("uso_excessivo_dois_pontos", False)
        est_trans = _esc(estilo.get("transicoes_entre_secoes", "—"))
        est_listas= estilo.get("listas_desnecessarias", False)
        est_obs   = _esc(estilo.get("observacoes") or "")

        def _check(v) -> str:
            if isinstance(v, bool):
                return '<span class="cross">✗</span>' if v else '<span class="check">✓</span>'
            return _esc(str(v))

        # Recomendações
        recs = analise.get("recomendacoes_priorizadas") or []

        # Coesão interna (redundâncias + fluxo geral→específico)
        redundancias = coesao.get("redundancias") or []
        sintese_coesao = _esc(coesao.get("sintese") or "")
        fluxo = _esc(coesao.get("fluxo_geral_especifico") or "—")
        obs_fluxo = _esc(coesao.get("observacoes_fluxo") or "")
        encadeamento = _esc(coesao.get("encadeamento") or "—")
        obs_encad = _esc(coesao.get("observacoes_encadeamento") or "")

        red_html = ""
        for red in redundancias:
            red_html += f"""<div class="par-item">
              <div class="argumento"><b>Trecho repetido:</b> {_esc(red.get("trecho_atual",""))}</div>
              <div class="fonte">🔁 já presente em: {_esc(red.get("capitulo_relacionado",""))} — {_esc(red.get("sugestao",""))}</div>
            </div>"""

        n_chars_fmt = f"{cap.get('n_chars', 0):,}"

        html += f"""
<div class="cap-card">
  <div class="cap-header" onclick="toggleCap(this)">
    <div>
      <div class="cap-title">{icon_pont} {titulo_cap}</div>
      <div style="font-size:.75rem;color:var(--cinza);margin-top:3px">{n_chars_fmt} caracteres</div>
    </div>
    <div style="display:flex;gap:8px;align-items:center">
      <span class="badge {badge_cls}">{_esc(pont or "—")}</span>
      <span style="color:var(--cinza);font-size:1.1rem">▸</span>
    </div>
  </div>
  <div class="cap-body">

    {"<div class='resumo-exec'>" + resumo_exec + "</div>" if resumo_exec else ""}

    {f'''<div class="bloco">
      <div class="bloco-titulo">⚠ Problemas Estruturais</div>
      {probs_html}
    </div>''' if probs_html else '<div class="bloco"><div class="bloco-titulo">✅ Sem problemas estruturais detectados</div></div>'}

    <div class="bloco" style="display:grid;grid-template-columns:1fr 1fr;gap:12px">
      <div>
        <div class="bloco-titulo">📚 Adequação Teórica: {teo_av}</div>
        {"<ul class='lista'>" + "".join(f"<li>✓ {_esc(p)}</li>" for p in teo_pos) + "</ul>" if teo_pos else ""}
        {"<div style='margin-top:6px'><b style='font-size:.78rem;color:var(--vermelho)'>Lacunas:</b><ul class='lista'>" + "".join(f"<li>{_esc(l)}</li>" for l in teo_lac) + "</ul></div>" if teo_lac else ""}
      </div>
      <div>
        <div class="bloco-titulo">🔗 Integração Empírica: {emp_av}</div>
        <p style="font-size:.85rem">{emp_obs}</p>
      </div>
    </div>

    <div class="bloco">
      <div class="bloco-titulo">✍ Estilo & Escrita</div>
      <table class="abnt">
        <tr><td>Parágrafos iniciando com autor</td><td>{_check(est_autor > 0)} {est_autor}</td></tr>
        <tr><td>Uso de travessão em prosa</td><td>{_check(est_trav)}</td></tr>
        <tr><td>Dois-pontos excessivos</td><td>{_check(est_2pts)}</td></tr>
        <tr><td>Transições entre seções</td><td>{est_trans}</td></tr>
        <tr><td>Listas desnecessárias</td><td>{_check(est_listas)}</td></tr>
      </table>
      {"<p style='font-size:.82rem;color:var(--cinza);margin-top:8px'>" + est_obs + "</p>" if est_obs else ""}
    </div>

    <div class="bloco">
      <div class="bloco-titulo">📏 Conformidade ABNT: {abnt_status}</div>
      {"<ul class='lista'>" + "".join(f"<li class='cross'>✗ {_esc(p)}</li>" for p in abnt_probs) + "</ul>" if abnt_probs else "<p style='font-size:.85rem;color:var(--verde)'>✓ Nenhum problema ABNT identificado</p>"}
    </div>

    {f'''<div class="bloco">
      <div class="bloco-titulo">📋 Recomendações Priorizadas</div>
      <ul class="lista">{"".join(f"<li>{_esc(r)}</li>" for r in recs)}</ul>
    </div>''' if recs else ""}

    <div class="bloco">
      <div class="bloco-titulo">🔗 Coesão Interna: fluxo geral→específico ({fluxo}) | encadeamento ({encadeamento})</div>
      {"<p style='font-size:.85rem'>" + obs_fluxo + "</p>" if obs_fluxo else ""}
      {"<p style='font-size:.85rem'>" + obs_encad + "</p>" if obs_encad else ""}
    </div>

    {f'''<div class="bloco">
      <div class="bloco-titulo">🔁 Redundâncias com outros capítulos</div>
      {red_html}
      {"<div class='sintese'>" + sintese_coesao + "</div>" if sintese_coesao else ""}
    </div>''' if (red_html or sintese_coesao) else ""}

  </div>
</div>
"""

    html += f"""
</main>
<footer>Análise gerada automaticamente em {_esc(ts)} · Framework genérico de dissertações</footer>
<script>
function toggleCap(header) {{
  const body = header.nextElementSibling;
  const arrow = header.querySelector('span:last-child');
  body.classList.toggle('aberto');
  arrow.textContent = body.classList.contains('aberto') ? '▾' : '▸';
}}
</script>
</body>
</html>"""

    (rel_dir / "relatorio_analise.html").write_text(html, encoding="utf-8")

    # ── TXT resumido ─────────────────────────────────────────────────────
    sep = "=" * 70
    L = [sep, "  ANÁLISE CRÍTICA DE DISSERTAÇÃO", f"  Gerado em: {ts}", sep, ""]
    L += [
        f"Autor       : {meta.get('autor') or '—'}",
        f"Título      : {meta.get('titulo') or '—'}",
        f"Instituição : {meta.get('instituicao') or '—'}",
        f"Orientador  : {meta.get('orientador') or '—'}",
        f"Programa    : {meta.get('programa') or '—'}",
        f"Ano         : {meta.get('ano') or '—'}",
        "",
        f"RESUMO: {n_aprovado} aprovados | {n_ressalvas} com ressalvas | {n_revisao} requerem revisão",
        f"Coesão interna: {total_red} redundância(s) entre capítulos identificada(s)",
        "",
    ]
    for cap in caps:
        an = cap.get("analise") or {}
        L += [
            f"{'─' * 60}",
            f"CAPÍTULO: {cap.get('titulo', '')}",
            f"Pontuação: {an.get('pontuacao_geral', '—')}",
        ]
        resumo_exec = an.get("resumo_executivo")
        if resumo_exec:
            L.append(f"Resumo: {resumo_exec}")
        prbs = an.get("problemas_estruturais") or []
        if prbs:
            L.append("Problemas:")
            for p in prbs:
                L.append(f"  [{p.get('gravidade', '?')}] {p.get('descricao', '')}")
        recs = an.get("recomendacoes_priorizadas") or []
        if recs:
            L.append("Recomendações:")
            for r in recs:
                L.append(f"  {r}")
        coesao = cap.get("coesao") or {}
        redundancias = coesao.get("redundancias") or []
        if redundancias:
            L.append(f"Redundâncias com outros capítulos: {len(redundancias)}")
            for r in redundancias:
                L.append(f"  - repete trecho já presente em: {r.get('capitulo_relacionado','')} — {r.get('sugestao','')}")
        if coesao.get("fluxo_geral_especifico"):
            L.append(f"Fluxo geral→específico: {coesao.get('fluxo_geral_especifico')}")
        if coesao.get("encadeamento"):
            L.append(f"Encadeamento com capítulos vizinhos: {coesao.get('encadeamento')}")
        L.append("")

    (rel_dir / "relatorio_analise.txt").write_text("\n".join(L), encoding="utf-8")


# ════════════════════════════════════════════════════════════════════════
# MAPA SEMÂNTICO (visão geral do documento — usado para evitar redundância
# e garantir "leitura completa" sem estourar o contexto)
# ════════════════════════════════════════════════════════════════════════

_PROMPT_MAPA = """Você é um assistente acadêmico criando um MAPA SEMÂNTICO de um capítulo de dissertação,
para ser usado depois como contexto ao revisar o documento inteiro (evitar redundâncias entre capítulos).

CAPÍTULO: «{titulo_cap}»

TEXTO (até {n_chars} caracteres):
\"\"\"
{texto_cap}
\"\"\"

Retorne SOMENTE um JSON válido (sem markdown, sem texto antes ou depois):
{{
  "tema_central": "1 frase resumindo do que trata o capítulo",
  "argumentos_chave": ["argumento 1", "argumento 2", "argumento 3"],
  "conceitos_recorrentes": ["conceito 1", "conceito 2"],
  "citacoes_principais": ["Autor, Ano: ideia citada"]
}}
"""


def gerar_mapa_semantico(capitulos: list[tuple[str, str]], api_key: str, log_fn=None) -> list[dict]:
    """
    Gera um resumo estruturado de cada capítulo (1 chamada por capítulo).
    Esse mapa é usado como contexto compacto representando o documento inteiro,
    permitindo detectar redundâncias entre capítulos distantes sem reenviar o texto todo.
    """
    client = anthropic.Anthropic(api_key=api_key)
    mapa: list[dict] = []
    for i, (titulo_cap, conteudo) in enumerate(capitulos, 1):
        if log_fn:
            log_fn(f"   🗺 Mapeando [{i}/{len(capitulos)}]: «{titulo_cap[:50]}»…")
        trecho = conteudo[:_MAX_CHARS_CAP]
        prompt = _PROMPT_MAPA.format(titulo_cap=titulo_cap, texto_cap=trecho, n_chars=len(trecho))
        try:
            resumo = _chamar_claude(client, prompt, max_tokens=1000, model="claude-haiku-4-5")
        except Exception as e:
            resumo = {"_erro": str(e)}
        resumo["titulo"] = titulo_cap
        mapa.append(resumo)
    return mapa


def _mapa_para_contexto(mapa: list[dict], excluir_idx: int = -1, max_chars: int = 6000) -> str:
    """Serializa o mapa semântico (exceto o capítulo atual) em texto compacto."""
    linhas = []
    for i, m in enumerate(mapa):
        if i == excluir_idx or m.get("_erro"):
            continue
        linhas.append(f"• {m.get('titulo','')}: {m.get('tema_central','')}")
        for arg in (m.get("argumentos_chave") or [])[:2]:
            linhas.append(f"   - {arg}")
        for cit in (m.get("citacoes_principais") or [])[:2]:
            linhas.append(f"   - cita: {cit}")
    texto = "\n".join(linhas)
    return texto[:max_chars]


# ════════════════════════════════════════════════════════════════════════
# REESCRITA DE CAPÍTULOS
# ════════════════════════════════════════════════════════════════════════

_PROMPT_REESCRITA = """Você é um revisor acadêmico sênior reescrevendo um capítulo de TCC/dissertação brasileira,
para gerar a VERSÃO FINAL CORRIGIDA do documento.

VISÃO GERAL DO RESTANTE DO DOCUMENTO (use para evitar repetir argumentos/conceitos já tratados em outros capítulos):
\"\"\"
{mapa_geral}
\"\"\"

CAPÍTULO ATUAL: «{titulo_cap}»

TEXTO ORIGINAL DO CAPÍTULO:
\"\"\"
{texto_cap}
\"\"\"

PROBLEMAS IDENTIFICADOS NA ANÁLISE CRÍTICA (corrija-os):
{problemas}

RECOMENDAÇÕES PRIORIZADAS:
{recomendacoes}

REDUNDÂNCIAS COM OUTROS CAPÍTULOS (remova ou resuma conforme indicado):
{consonancia}

INSTRUÇÕES PARA A REESCRITA:
1. Mantenha a extensão e profundidade compatíveis com um TCC/dissertação — NÃO resuma demais, é para ser uma versão completa e revisada, não um resumo.
2. Remova redundâncias: frases ou argumentos repetidos dentro do próprio capítulo.
3. Remova ou sinalize argumentos que já aparecem em outros capítulos (conforme a visão geral e a lista de redundâncias acima) — referencie a seção correspondente em vez de repetir.
4. Reorganize a estrutura interna do capítulo do geral para o específico (contexto amplo → discussão específica → conclusão do capítulo).
5. Corrija os problemas estruturais e ABNT identificados.
6. Preserve citações ABNT válidas (autor, ano) e dados/resultados empíricos.
7. Escreva em português formal acadêmico, terceira pessoa/impessoal.

Retorne SOMENTE o texto revisado do capítulo (sem comentários, sem markdown, sem JSON — apenas o texto corrido do capítulo, incluindo o título)."""


def reescrever_capitulo(
    titulo_cap: str,
    texto_cap: str,
    mapa_geral: str,
    analise: dict,
    coesao: dict,
    api_key: str,
) -> str:
    client = anthropic.Anthropic(api_key=api_key)

    probs = analise.get("problemas_estruturais") or []
    problemas_str = "\n".join(
        f"- [{p.get('gravidade','?')}] {p.get('descricao','')}" for p in probs
    ) or "Nenhum problema estrutural relevante identificado."

    recs = analise.get("recomendacoes_priorizadas") or []
    recs_str = "\n".join(f"- {r}" for r in recs) or "Nenhuma recomendação adicional."

    redundancias = coesao.get("redundancias") or []
    if redundancias:
        consonancia_str = "\n".join(
            f"- Trecho: \"{r.get('trecho_atual','')}\" → "
            f"já presente em: {r.get('capitulo_relacionado','')} — {r.get('sugestao','')}"
            for r in redundancias
        )
    else:
        consonancia_str = "Nenhuma redundância identificada para este capítulo."

    prompt = _PROMPT_REESCRITA.format(
        mapa_geral=mapa_geral or "(documento de capítulo único)",
        titulo_cap=titulo_cap,
        texto_cap=texto_cap[:_MAX_CHARS_CAP],
        problemas=problemas_str,
        recomendacoes=recs_str,
        consonancia=consonancia_str,
    )

    resp = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.content[0].text.strip()


_PROMPT_REVISAO_TRANSICAO = """Você é um revisor acadêmico fazendo a REVISÃO FINAL de coesão entre capítulos consecutivos
de um TCC/dissertação já reescrito.

FIM DO CAPÍTULO ANTERIOR («{titulo_ant}»):
\"\"\"
{fim_ant}
\"\"\"

CAPÍTULO ATUAL («{titulo_atual}»), TEXTO COMPLETO:
\"\"\"
{texto_atual}
\"\"\"

INÍCIO DO PRÓXIMO CAPÍTULO («{titulo_prox}»):
\"\"\"
{inicio_prox}
\"\"\"

Revise APENAS o capítulo atual para:
1. Garantir transição fluida com o capítulo anterior e o próximo (sem repetir literalmente o que já foi dito).
2. Remover qualquer frase ou parágrafo redundante que se repita quase identicamente em relação aos trechos mostrados acima.
3. Manter todo o restante do conteúdo e extensão.

Retorne SOMENTE o texto revisado do capítulo atual (sem comentários, sem markdown)."""


def revisar_transicoes(capitulos_revisados: list[dict], api_key: str, log_fn=None) -> list[dict]:
    """
    Passada final: revisa cada capítulo considerando o fim do anterior e o início do próximo,
    para garantir coesão geral e remover repetições residuais entre capítulos adjacentes.
    """
    client = anthropic.Anthropic(api_key=api_key)
    n = len(capitulos_revisados)
    revisados_final = []

    for i, cap in enumerate(capitulos_revisados):
        titulo = cap["titulo"]
        texto = cap["texto"]

        if n == 1:
            revisados_final.append(cap)
            continue

        fim_ant = capitulos_revisados[i-1]["texto"][-1500:] if i > 0 else "(é o primeiro capítulo)"
        titulo_ant = capitulos_revisados[i-1]["titulo"] if i > 0 else "—"
        inicio_prox = capitulos_revisados[i+1]["texto"][:1500] if i < n-1 else "(é o último capítulo)"
        titulo_prox = capitulos_revisados[i+1]["titulo"] if i < n-1 else "—"

        if log_fn:
            log_fn(f"   📝 Revisão final [{i+1}/{n}]: «{titulo[:50]}»…")

        prompt = _PROMPT_REVISAO_TRANSICAO.format(
            titulo_ant=titulo_ant, fim_ant=fim_ant,
            titulo_atual=titulo, texto_atual=texto[:_MAX_CHARS_CAP],
            titulo_prox=titulo_prox, inicio_prox=inicio_prox,
        )
        try:
            resp = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}],
            )
            texto_final = resp.content[0].text.strip()
        except Exception as e:
            if log_fn:
                log_fn(f"   ⚠ Erro na revisão final do capítulo: {e}", "warn")
            texto_final = texto

        revisados_final.append({"titulo": titulo, "texto": texto_final})

    return revisados_final


# ════════════════════════════════════════════════════════════════════════
# EXPORTAÇÃO DO .DOCX REVISADO
# ════════════════════════════════════════════════════════════════════════

def gerar_docx_revisado(metadados: dict, capitulos_revisados: list[dict], output_path: Path) -> Path:
    """Monta o documento .docx final a partir dos capítulos reescritos."""
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()

    # Página de rosto
    titulo = doc.add_heading(metadados.get("titulo") or "Dissertação Revisada", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for campo, label in [
        ("autor", "Autor"), ("instituicao", "Instituição"), ("programa", "Programa"),
        ("orientador", "Orientador"), ("area", "Área"), ("ano", "Ano"),
    ]:
        valor = metadados.get(campo)
        if valor:
            info.add_run(f"{label}: {valor}\n")

    aviso = doc.add_paragraph()
    aviso_run = aviso.add_run(
        "Documento gerado automaticamente — versão revisada com base em análise crítica "
        "de IA, removendo redundâncias e corrigindo inconsistências de consonância com as fontes citadas. "
        "Recomenda-se revisão final pelo autor antes da entrega."
    )
    aviso_run.italic = True
    aviso_run.font.size = Pt(9)

    doc.add_page_break()

    # Capítulos
    for cap in capitulos_revisados:
        texto = cap["texto"]
        linhas = [l for l in texto.split("\n") if l.strip()]
        if not linhas:
            continue

        # Primeira linha não-vazia = título do capítulo
        primeiro = linhas[0].strip()
        resto = linhas[1:]

        # Heuristica: titulo curto -> Heading; senao usa o titulo original
        if len(primeiro) <= 120:
            doc.add_heading(primeiro, level=1)
        else:
            doc.add_heading(cap["titulo"], level=1)
            resto = linhas

        for par in resto:
            par = par.strip()
            if not par:
                continue
            p = doc.add_paragraph(par)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path


# ════════════════════════════════════════════════════════════════════════
# ORQUESTRADOR — DISSERTAÇÃO REVISADA
# ════════════════════════════════════════════════════════════════════════

def gerar_dissertacao_revisada(
    capitulos: list[tuple[str, str]],
    resultados_caps: list[dict],
    metadados: dict,
    api_key: str,
    output_path: Path,
    log_fn,
) -> Path:
    """
    Pipeline completo de geração da dissertação revisada:
      1. Mapa semântico do documento (visão geral, evita redundância)
      2. Reescrita capítulo a capítulo (corrige problemas, consonância, redundância, ordem geral->especifico)
      3. Revisão final de transições/coesão entre capítulos
      4. Exporta .docx
    """
    log_fn("🗺 ETAPA 6: Gerando mapa semântico do documento completo…", "etapa")
    mapa = gerar_mapa_semantico(capitulos, api_key, log_fn=log_fn)
    log_fn(f"   Mapa semântico gerado para {len(mapa)} capítulo(s)")

    log_fn("✍ ETAPA 7: Reescrevendo capítulos (corrigindo redundâncias, consonância e estrutura)…", "etapa")
    capitulos_revisados: list[dict] = []
    n = len(capitulos)
    for i, (titulo_cap, conteudo) in enumerate(capitulos):
        log_fn(f"   [{i+1}/{n}] Reescrevendo: «{titulo_cap[:50]}»…")
        analise = (resultados_caps[i] or {}).get("analise") or {}
        coesao = (resultados_caps[i] or {}).get("coesao") or {}
        mapa_contexto = _mapa_para_contexto(mapa, excluir_idx=i)
        try:
            texto_revisado = reescrever_capitulo(titulo_cap, conteudo, mapa_contexto, analise, coesao, api_key)
        except Exception as e:
            log_fn(f"   ⚠ Erro ao reescrever capítulo, mantendo original: {e}", "warn")
            texto_revisado = f"{titulo_cap}\n\n{conteudo}"
        capitulos_revisados.append({"titulo": titulo_cap, "texto": texto_revisado})

    log_fn("📝 ETAPA 8: Revisão final de coesão entre capítulos…", "etapa")
    capitulos_revisados = revisar_transicoes(capitulos_revisados, api_key, log_fn=log_fn)

    log_fn("💾 ETAPA 9: Gerando arquivo .docx revisado…", "etapa")
    caminho = gerar_docx_revisado(metadados, capitulos_revisados, output_path)
    log_fn(f"   ✓ Dissertação revisada salva: {caminho.name}", "ok")

    return caminho
