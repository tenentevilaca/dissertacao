"""
Gera o documento de dissertação atualizado (.docx) e o relatório de análise (.docx).
"""

import json
from datetime import datetime
from pathlib import Path

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers de formatação ────────────────────────────────────────────────

def _set_run_cor(run, r, g, b):
    run.font.color.rgb = RGBColor(r, g, b)


def _add_heading(doc: Document, texto: str, nivel: int):
    para = doc.add_heading(texto, level=nivel)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para


def _add_para(doc: Document, texto: str, bold=False, italic=False, cor=None):
    para = doc.add_paragraph()
    run = para.add_run(texto)
    run.bold = bold
    run.italic = italic
    if cor:
        _set_run_cor(run, *cor)
    para.paragraph_format.space_after = Pt(6)
    return para


def _add_box(doc: Document, titulo: str, conteudo: str, cor_titulo=(30, 80, 160)):
    """Adiciona uma caixa de texto estilizada (usando parágrafo sombreado)."""
    p = doc.add_paragraph()
    run = p.add_run(f"▌ {titulo}")
    run.bold = True
    _set_run_cor(run, *cor_titulo)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    for linha in conteudo.split("\n"):
        if linha.strip():
            lp = doc.add_paragraph(f"    {linha.strip()}")
            lp.paragraph_format.left_indent = Inches(0.3)
            lp.paragraph_format.space_after = Pt(2)


# ── Relatório de análise ─────────────────────────────────────────────────

def gerar_relatorio(
    metadados: dict,
    resultados_capitulos: list[dict],
    caminho_saida: Path,
) -> Path:
    """Gera relatorio_analise.docx com diagnóstico completo."""
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Capa
    doc.add_heading("RELATÓRIO DE ANÁLISE DE DISSERTAÇÃO", level=0)
    doc.add_paragraph()
    _add_para(doc, f"Autor: {metadados.get('autor', 'N/D')}", bold=True)
    _add_para(doc, f"Título: {metadados.get('titulo', 'N/D')}", bold=True)
    _add_para(doc, f"Programa: {metadados.get('programa', 'N/D')} — {metadados.get('instituicao', 'N/D')}")
    _add_para(doc, f"Metodologia: {metadados.get('metodologia', 'N/D')}")
    _add_para(doc, f"Data da análise: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_page_break()

    # Sumário executivo
    _add_heading(doc, "SUMÁRIO EXECUTIVO", level=1)
    pontuacoes = [
        r["diagnostico"].get("pontuacao_geral", 0)
        for r in resultados_capitulos
        if isinstance(r.get("diagnostico"), dict)
    ]
    media = sum(pontuacoes) / len(pontuacoes) if pontuacoes else 0
    _add_para(doc, f"Pontuação média geral: {media:.1f}/10")
    _add_para(doc, f"Capítulos analisados: {len(resultados_capitulos)}")
    doc.add_paragraph()

    # Análise por capítulo
    for res in resultados_capitulos:
        if res["numero"] == 0:
            continue  # pula pré-textual

        _add_heading(doc, f"CAPÍTULO {res['numero']} — {res['titulo']}", level=1)
        diag = res.get("diagnostico", {})

        if not isinstance(diag, dict):
            _add_para(doc, str(diag))
            continue

        pontos = diag.get("pontuacao_geral", "N/D")
        _add_para(doc, f"Pontuação: {pontos}/10", bold=True)

        resumo = diag.get("resumo_diagnostico", "")
        if resumo:
            _add_para(doc, resumo, italic=True)

        doc.add_paragraph()

        # Problemas
        for categoria, label, cor in [
            ("problemas_estruturais", "Problemas Estruturais", (180, 20, 20)),
            ("problemas_teoricos", "Adequação Teórica", (140, 60, 20)),
            ("problemas_abnt", "Conformidade ABNT", (20, 80, 140)),
            ("problemas_integracao", "Integração Teórico-Empírica", (80, 20, 140)),
            ("problemas_coesao", "Coesão e Coerência", (20, 120, 80)),
        ]:
            itens = diag.get(categoria, [])
            if itens:
                _add_box(doc, label, "\n".join(f"• {i}" for i in itens), cor_titulo=cor)

        positivos = diag.get("pontos_positivos", [])
        if positivos:
            _add_box(doc, "Pontos Positivos", "\n".join(f"✓ {i}" for i in positivos), cor_titulo=(30, 120, 30))

        lacunas = diag.get("lacunas", [])
        if lacunas:
            _add_box(doc, "Lacunas Identificadas", "\n".join(f"△ {i}" for i in lacunas), cor_titulo=(160, 100, 20))

        recomendacoes = diag.get("recomendacoes") if isinstance(diag, dict) else []
        recomendacoes = res.get("diagnostico", {}).get("recomendacoes", []) if isinstance(res.get("diagnostico"), dict) else []
        if recomendacoes:
            _add_box(doc, "Recomendações Aplicadas na Reescrita", "\n".join(f"→ {r}" for r in recomendacoes), cor_titulo=(20, 80, 160))

        doc.add_paragraph()

    saida = caminho_saida / "relatorio_analise.docx"
    doc.save(str(saida))
    return saida


# ── Dissertação reescrita ────────────────────────────────────────────────

def gerar_dissertacao_reescrita(
    caminho_original: Path,
    metadados: dict,
    resultados_capitulos: list[dict],
    caminho_saida: Path,
) -> Path:
    """
    Gera dissertacao_reescrita.docx combinando:
    - elementos pré-textuais originais (capa, folha de rosto, sumário, etc.)
    - capítulos reescritos
    Preserva a formatação base do original.
    """
    # Tenta abrir o original para extrair estrutura/estilos
    try:
        doc_orig = Document(str(caminho_original))
        doc = _clonar_estilos(doc_orig)
    except Exception:
        doc = Document()
        _configurar_estilos_padrao(doc)

    # Cabeçalho
    _add_heading(doc, metadados.get("titulo", "Dissertação"), level=0)
    _add_para(doc, f"Autor: {metadados.get('autor', '')}", bold=True)
    _add_para(doc, f"Programa: {metadados.get('programa', '')} — {metadados.get('instituicao', '')}")
    _add_para(doc, f"Orientador: {metadados.get('orientador', '')}")
    _add_para(doc, f"[Versão revisada — {datetime.now().strftime('%d/%m/%Y')}]", italic=True, cor=(100, 100, 100))
    doc.add_page_break()

    # Capítulos reescritos
    for res in resultados_capitulos:
        if res["numero"] == 0:
            # Elementos pré-textuais: copia do original
            for linha in res.get("texto_original", "").split("\n"):
                if linha.strip():
                    doc.add_paragraph(linha.strip())
            doc.add_page_break()
            continue

        texto_reescrito = res.get("texto_reescrito", "") or res.get("texto_original", "")
        _inserir_capitulo_reescrito(doc, res["numero"], res["titulo"], texto_reescrito)
        doc.add_page_break()

    saida = caminho_saida / "dissertacao_reescrita.docx"
    doc.save(str(saida))
    return saida


def _inserir_capitulo_reescrito(doc: Document, numero: int, titulo: str, texto: str):
    """Insere um capítulo reescrito no documento."""
    linhas = texto.split("\n")
    primeiro = True

    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue

        # Detecta se a linha é título/subtítulo
        if _parece_titulo(linha):
            nivel = _nivel_titulo(linha)
            if primeiro:
                _add_heading(doc, linha, level=1)
                primeiro = False
            else:
                _add_heading(doc, linha, level=nivel)
        else:
            p = doc.add_paragraph(linha)
            p.paragraph_format.space_after = Pt(6)
            if primeiro:
                primeiro = False


def _parece_titulo(texto: str) -> bool:
    """Heurística simples: curto, sem ponto final, todo caps ou numerado."""
    if len(texto) > 150:
        return False
    if texto.endswith(".") or texto.endswith(","):
        return False
    if texto.isupper() and len(texto) > 3:
        return True
    import re
    if re.match(r"^\d+[\.\d]*\s+[A-ZÁÉÍÓÚ]", texto):
        return True
    return False


def _nivel_titulo(texto: str) -> int:
    import re
    m = re.match(r"^(\d+)(\.\d+)*", texto)
    if not m:
        return 1
    partes = texto[:m.end()].split(".")
    return min(len(partes), 3)


def _clonar_estilos(doc_orig: Document) -> Document:
    """Cria novo documento copiando os estilos do original."""
    novo = Document()
    # Copia fonte base
    try:
        estilo = doc_orig.styles["Normal"]
        novo_estilo = novo.styles["Normal"]
        novo_estilo.font.name = estilo.font.name or "Arial"
        novo_estilo.font.size = estilo.font.size or Pt(12)
    except Exception:
        pass
    return novo


def _configurar_estilos_padrao(doc: Document):
    try:
        estilo = doc.styles["Normal"]
        estilo.font.name = "Arial"
        estilo.font.size = Pt(12)
    except Exception:
        pass
